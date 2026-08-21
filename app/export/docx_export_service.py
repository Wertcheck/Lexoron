"""DraftDocxExportService – Export EINES Entwurfs als formatierte `.docx`
(Schriftsatz-Generator, 20.08.; Briefkopf-/Signatur-Verwaltung, Nachtrag
20.08.).

Getrennt von `MatterExportService` (app/export/service.py, ZIP-Export der
GESAMTEN Akte für Auskunftsersuchen/Archivierung) - dieser Service liefert
gezielt EIN druckfertiges Word-Dokument für genau eine Entwurfsversion.

Briefkopf-/Signatur-Aufbau (Logo/Anschrift/Unterschrift) lebt seit dem
Dokumentengenerator (Block 3, 20.08.) in app/export/letterhead.py - EINE
gemeinsame Implementierung statt einer zweiten, fast identischen Kopie in
app/document_generator/docx_export.py. Optional (`firm_profile=None` oder
ein leerer Datensatz möglich): solange auf der Kanzlei-Profilseite
(/dashboard/settings/profile) kein Kanzleiname eingetragen wurde, bleibt
der Export ohne Briefkopf (ehrlicher als ein Platzhalter-Absender).
"""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from docx.shared import Pt

from app.export.letterhead import add_signature_block, build_header, has_letterhead_content, has_signature_content
from app.models import Draft, FirmProfile, Matter

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DraftDocxExportService:
    def export_draft(
        self, draft: Draft, matter: Matter, firm_profile: FirmProfile | None = None
    ) -> BytesIO:
        document = DocxDocument()

        style = document.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        if has_letterhead_content(firm_profile):
            build_header(document, firm_profile)

        document.add_heading(matter.title or "Schriftsatz", level=1)
        meta = document.add_paragraph()
        meta.add_run(
            f"Entwurf Version {draft.version} · Stand "
            f"{draft.updated_at.strftime('%d.%m.%Y')}"
        ).italic = True

        # Leerzeilen als Absatzgrenzen - der Entwurfstext selbst ist reiner
        # Fließtext ohne eigene Formatierungssyntax (siehe Draft.content).
        for block in draft.content.split("\n\n"):
            block = block.strip()
            if block:
                document.add_paragraph(block)

        if has_signature_content(firm_profile):
            add_signature_block(document, firm_profile)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer
