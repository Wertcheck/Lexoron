"""Wiederverwendbarer Briefkopf/Signatur-Aufbau für Word-Exporte.

Extrahiert (20.08., Dokumentengenerator/Block 3) aus
`DraftDocxExportService` (app/export/docx_export_service.py) - GENAU EINE
Implementierung der OXML-Rahmenlinie/Kopf-/Signaturlogik statt einer
zweiten, fast identischen Kopie für `GeneratedDocumentDocxExportService`
(app/document_generator/docx_export.py). Verhalten unverändert gegenüber
vorher (siehe tests/test_draft_docx_export.py) - reine Verschiebung, keine
Änderung der Ausgabe."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.models import FirmProfile

_DIVIDER_COLOR = "CBD5E1"  # entspricht --paper-line (#e2e8f0-nah, druckfertig etwas kraeftiger)
_LOGO_HEIGHT_CM = 1.6
_SIGNATURE_HEIGHT_CM = 2.2


def image_exists(path: str | None) -> bool:
    return bool(path) and Path(path).exists()


def has_letterhead_content(firm_profile: FirmProfile | None) -> bool:
    return firm_profile is not None and (
        firm_profile.firm_name.strip() or image_exists(firm_profile.logo_path)
    )


def has_signature_content(firm_profile: FirmProfile | None) -> bool:
    return firm_profile is not None and (
        image_exists(firm_profile.signature_path) or (firm_profile.signatory_name or "").strip()
    )


def add_bottom_border(paragraph, *, size: int = 6, space: int = 6) -> None:
    """Fügt dem Absatz eine dünne untere Rahmenlinie hinzu (python-docx hat
    dafür keine High-Level-API - direkte OXML-Manipulation, Standardmuster
    für Word-Absatzrahmen). `size` in 1/8pt (6 = 0.75pt, eine dezente
    Trennlinie statt eines dicken Balkens)."""
    paragraph_format_element = paragraph._p.get_or_add_pPr()
    border_element = OxmlElement("w:pBdr")
    bottom_element = OxmlElement("w:bottom")
    bottom_element.set(qn("w:val"), "single")
    bottom_element.set(qn("w:sz"), str(size))
    bottom_element.set(qn("w:space"), str(space))
    bottom_element.set(qn("w:color"), _DIVIDER_COLOR)
    border_element.append(bottom_element)
    paragraph_format_element.append(border_element)


def _address_and_contact_lines(firm_profile: FirmProfile) -> list[str]:
    address_line = ", ".join(
        part
        for part in (
            firm_profile.street,
            " ".join(p for p in (firm_profile.postal_code, firm_profile.city) if p) or None,
        )
        if part
    )
    contact_line = " · ".join(
        part for part in (firm_profile.phone, firm_profile.email, firm_profile.website) if part
    )
    return [line for line in (address_line, contact_line) if line]


def build_header(document: DocxDocument, firm_profile: FirmProfile) -> None:
    """Seiten-Kopfbereich: Logo (falls vorhanden) zentriert oben, darunter
    Kanzleiname/Anschrift/Kontakt, abgeschlossen mit einer Trennlinie.
    Läuft auf JEDER Seite des Dokuments (echter Word-Header, nicht nur ein
    vorangestellter Body-Absatz)."""
    header = document.sections[0].header
    header.is_linked_to_previous = False

    # Der frische Header hat bereits einen leeren Absatz - wiederverwenden
    # statt einen weiteren leeren voranzustellen.
    paragraphs_used = 0
    if image_exists(firm_profile.logo_path):
        logo_paragraph = header.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.add_run().add_picture(firm_profile.logo_path, height=Cm(_LOGO_HEIGHT_CM))
        paragraphs_used += 1

    firm_name = firm_profile.firm_name.strip()
    if firm_name:
        name_paragraph = header.add_paragraph() if paragraphs_used else header.paragraphs[0]
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_paragraph.add_run(firm_name)
        name_run.bold = True
        name_run.font.size = Pt(12)
        paragraphs_used += 1

    detail_lines = _address_and_contact_lines(firm_profile)
    last_paragraph = None
    for line in detail_lines:
        detail_paragraph = header.add_paragraph() if paragraphs_used else header.paragraphs[0]
        detail_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        detail_run = detail_paragraph.add_run(line)
        detail_run.font.size = Pt(9)
        paragraphs_used += 1
        last_paragraph = detail_paragraph

    # Trennlinie unter dem letzten tatsächlich befüllten Absatz - fällt
    # keiner der obigen Zweige zu (nur Logo, kein Text), landet sie unter
    # dem Logo-Absatz.
    add_bottom_border(last_paragraph or header.paragraphs[0])


def add_signature_block(document: DocxDocument, firm_profile: FirmProfile) -> None:
    """Unterschriften-Grafik + getippter Name als letzter Abschnitt des
    Dokumentinhalts (bewusst NICHT im Word-Footer, der auf jeder Seite
    wiederholt würde - eine Unterschrift gehört einmalig ans Ende)."""
    document.add_paragraph()  # Abstand zum Brieftext

    if image_exists(firm_profile.signature_path):
        signature_paragraph = document.add_paragraph()
        signature_paragraph.add_run().add_picture(
            firm_profile.signature_path, height=Cm(_SIGNATURE_HEIGHT_CM)
        )

    signatory_name = (firm_profile.signatory_name or "").strip()
    if signatory_name:
        name_paragraph = document.add_paragraph()
        name_paragraph.add_run(signatory_name)
