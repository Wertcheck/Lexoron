"""GeneratedDocumentDocxExportService – Export EINES generierten Dokuments
als formatierte `.docx` (Block 3, 20.08.).

Nutzt denselben Briefkopf-/Signaturaufbau wie der Schriftsatz-Generator
(app/export/letterhead.py, ursprünglich aus
app/export/docx_export_service.py extrahiert) - EIN gemeinsamer, bereits
getesteter Baustein statt einer zweiten Implementierung."""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from docx.shared import Pt

from app.export.letterhead import add_signature_block, build_header, has_letterhead_content, has_signature_content
from app.models import FirmProfile, GeneratedDocument, Matter

DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class GeneratedDocumentDocxExportService:
    def export(
        self,
        document: GeneratedDocument,
        matter: Matter,
        firm_profile: FirmProfile | None = None,
    ) -> BytesIO:
        docx_document = DocxDocument()

        style = docx_document.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        if has_letterhead_content(firm_profile):
            build_header(docx_document, firm_profile)

        docx_document.add_heading(document.title, level=1)
        meta = docx_document.add_paragraph()
        meta.add_run(
            f"Akte: {matter.title} · Erstellt {document.created_at.strftime('%d.%m.%Y')}"
        ).italic = True

        for block in document.content.split("\n\n"):
            block = block.strip()
            if block:
                docx_document.add_paragraph(block)

        if has_signature_content(firm_profile):
            add_signature_block(docx_document, firm_profile)

        buffer = BytesIO()
        docx_document.save(buffer)
        buffer.seek(0)
        return buffer
