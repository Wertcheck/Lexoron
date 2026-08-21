"""Tests für app/document_generator/docx_export.py und pdf_export.py
(Block 3, 20.08.)."""

from __future__ import annotations

from collections.abc import Iterator

import pytest
from docx import Document as DocxDocument
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.document_generator.docx_export import GeneratedDocumentDocxExportService
from app.document_generator.pdf_export import GeneratedDocumentPdfExportService
from app.document_generator.service import generate_from_template
from app.models import Client, DocumentTemplate, FirmProfile, Matter
from app.models.base import Base


@pytest.fixture()
def db_session() -> Iterator[Session]:
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    TestingSessionLocal = sessionmaker(bind=engine)
    session = TestingSessionLocal()
    try:
        yield session
    finally:
        session.close()
        engine.dispose()


def _generated_document(db: Session, content: str = "Sehr geehrte Damen und Herren,\n\nTestinhalt."):
    client = Client(name="Muster GmbH", client_number="M-1")
    db.add(client)
    db.flush()
    matter = Matter(client_id=client.id, title="Testakte", status="open")
    db.add(matter)
    db.flush()
    template = DocumentTemplate(name="Testvorlage", content=content, version=1)
    db.add(template)
    db.commit()
    result = generate_from_template(db, template, matter, actor="anwalt@kanzlei.test")
    return result.document, matter


# --- DOCX ---


def test_docx_export_without_firm_profile_produces_valid_document(db_session: Session) -> None:
    document, matter = _generated_document(db_session)
    buffer = GeneratedDocumentDocxExportService().export(document, matter, firm_profile=None)
    docx = DocxDocument(buffer)
    full_text = "\n".join(p.text for p in docx.paragraphs)
    assert document.title in full_text
    assert "Testinhalt." in full_text


def test_docx_export_with_firm_profile_includes_letterhead(db_session: Session) -> None:
    document, matter = _generated_document(db_session)
    firm_profile = FirmProfile(firm_name="Kanzlei Muster & Partner")
    buffer = GeneratedDocumentDocxExportService().export(document, matter, firm_profile=firm_profile)
    docx = DocxDocument(buffer)
    header_text = "\n".join(p.text for p in docx.sections[0].header.paragraphs)
    assert "Kanzlei Muster & Partner" in header_text


def test_docx_export_reuses_shared_letterhead_module(db_session: Session) -> None:
    """Regressionsschutz: stellt sicher, dass der Export weiterhin
    app/export/letterhead.py nutzt (Wiederverwendung statt Duplikat, siehe
    Moduldocstring von docx_export.py)."""
    import app.document_generator.docx_export as module

    assert hasattr(module, "build_header")
    assert module.build_header.__module__ == "app.export.letterhead"


# --- PDF ---


def test_pdf_export_produces_valid_pdf_bytes(db_session: Session) -> None:
    document, matter = _generated_document(db_session)
    buffer = GeneratedDocumentPdfExportService().export(document, matter)
    content = buffer.read()
    assert content[:5] == b"%PDF-"
    assert len(content) > 100


def test_pdf_export_paginates_long_content_into_multiple_pages(db_session: Session) -> None:
    """Deterministische Paginierung darf bei sehr langem Text nicht
    haengen/abstuerzen (siehe Moduldocstring pdf_export.py)."""
    import pymupdf

    long_content = "\n\n".join(f"Absatz Nummer {i} mit etwas Text zum Füllen der Seite." * 3 for i in range(60))
    document, matter = _generated_document(db_session, content=long_content)
    buffer = GeneratedDocumentPdfExportService().export(document, matter)

    pdf = pymupdf.open(stream=buffer.read(), filetype="pdf")
    assert pdf.page_count > 1
    pdf.close()


def test_pdf_export_handles_single_pathologically_long_paragraph(db_session: Session) -> None:
    """Ein einzelner, extrem langer Absatz (kein doppeltes Zeilenumbruch-
    Zeichen) darf keine Endlosschleife ausloesen - siehe pdf_export.py:
    textwrap-basierte Zeilenumbrueche statt insert_textbox-Retry-Schleife."""
    import pymupdf

    huge_paragraph = "Wort " * 5000
    document, matter = _generated_document(db_session, content=huge_paragraph)
    buffer = GeneratedDocumentPdfExportService().export(document, matter)

    pdf = pymupdf.open(stream=buffer.read(), filetype="pdf")
    assert pdf.page_count > 1
    pdf.close()
