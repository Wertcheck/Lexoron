"""Tests für app/export/docx_export_service.py und die zugehörige Route
GET /dashboard/drafts/{draft_id}/export.docx (Schriftsatz-Generator, 20.08.;
Briefkopf-/Signatur-Verwaltung mit Logo/Unterschrift, Nachtrag 20.08.).
"""

from __future__ import annotations

import base64
from collections.abc import Iterator
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.db.session import get_db
from app.export.docx_export_service import DOCX_MEDIA_TYPE, DraftDocxExportService
from app.main import app
from app.models import AuditEvent, Client, Draft, FirmProfile, Matter
from app.models.base import Base
from tests.auth_test_utils import login_as_admin

# Kleinstmögliches valides PNG (1x1, transparent) - genug, damit
# python-docx `add_picture` es tatsächlich einbetten kann, ohne eine echte
# Bilddatei ins Repo aufnehmen zu müssen.
_TINY_PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUB"
    "AScY42YAAAAASUVORK5CYII="
)


def _write_tiny_png(path: Path) -> str:
    path.write_bytes(base64.b64decode(_TINY_PNG_BASE64))
    return str(path)


@pytest.fixture()
def db_session() -> Iterator[Session]:
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    TestingSessionLocal = sessionmaker(bind=engine)
    session = TestingSessionLocal()
    try:
        yield session
    finally:
        session.close()
        engine.dispose()


def _seed_draft(db: Session) -> Draft:
    client = Client(name="Testmandant GmbH")
    matter = Matter(client=client, title="Einspruch Steuerbescheid 2025")
    db.add_all([client, matter])
    db.commit()
    draft = Draft(matter_id=matter.id, content="Erster Absatz.\n\nZweiter Absatz.")
    db.add(draft)
    db.commit()
    return draft


def test_export_service_produces_readable_docx(db_session: Session) -> None:
    draft = _seed_draft(db_session)

    buffer = DraftDocxExportService().export_draft(draft, draft.matter)

    document = DocxDocument(BytesIO(buffer.getvalue()))
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "Einspruch Steuerbescheid 2025" in full_text
    assert "Erster Absatz." in full_text
    assert "Zweiter Absatz." in full_text


def test_export_service_without_firm_profile_has_no_letterhead(db_session: Session) -> None:
    draft = _seed_draft(db_session)

    buffer = DraftDocxExportService().export_draft(draft, draft.matter, firm_profile=None)

    document = DocxDocument(BytesIO(buffer.getvalue()))
    full_text = "\n".join(p.text for p in document.paragraphs)
    assert "Kanzlei" not in full_text.split("Einspruch")[0]


def test_export_service_with_empty_firm_profile_has_no_letterhead(db_session: Session) -> None:
    """Ein automatisch angelegter, aber noch nicht ausgefüllter
    FirmProfile (leerer firm_name, kein Logo) darf NICHT als Briefkopf
    erscheinen - siehe app/export/docx_export_service.py Moduldocstring."""
    draft = _seed_draft(db_session)
    empty_profile = FirmProfile(firm_name="")

    buffer = DraftDocxExportService().export_draft(draft, draft.matter, empty_profile)

    document = DocxDocument(BytesIO(buffer.getvalue()))
    assert document.paragraphs[0].text != ""  # keine leere Briefkopf-Zeile vorangestellt
    header_text = "".join(p.text for p in document.sections[0].header.paragraphs)
    assert header_text == ""


def test_export_service_with_firm_profile_adds_header_letterhead(db_session: Session) -> None:
    """Briefkopf landet im echten Word-Seitenkopf (erscheint auf jeder
    Seite), NICHT als vorangestellter Body-Absatz - siehe Moduldocstring."""
    draft = _seed_draft(db_session)
    profile = FirmProfile(
        firm_name="Kanzlei Mustermann Rechtsanwälte",
        street="Musterstraße 12",
        postal_code="10115",
        city="Berlin",
        phone="+49 30 1234567",
        email="kanzlei@beispiel.de",
        website=None,
    )

    buffer = DraftDocxExportService().export_draft(draft, draft.matter, profile)

    document = DocxDocument(BytesIO(buffer.getvalue()))
    header_text = "\n".join(p.text for p in document.sections[0].header.paragraphs)
    assert "Kanzlei Mustermann Rechtsanwälte" in header_text
    assert "Musterstraße 12, 10115 Berlin" in header_text
    assert "+49 30 1234567 · kanzlei@beispiel.de" in header_text
    # Briefkopf-Text darf NICHT im Fließtext des Dokuments auftauchen.
    body_text = "\n".join(p.text for p in document.paragraphs)
    assert "Kanzlei Mustermann Rechtsanwälte" not in body_text


def test_export_service_embeds_logo_in_header(db_session: Session, tmp_path: Path) -> None:
    draft = _seed_draft(db_session)
    logo_path = _write_tiny_png(tmp_path / "logo.png")
    profile = FirmProfile(firm_name="Kanzlei Mustermann Rechtsanwälte", logo_path=logo_path)

    buffer = DraftDocxExportService().export_draft(draft, draft.matter, profile)

    document = DocxDocument(BytesIO(buffer.getvalue()))
    # Bilder im Header liegen in einem eigenen OOXML-Part (header1.xml),
    # NICHT in document.inline_shapes (das deckt nur den Hauptteil ab) -
    # `HeaderPart` selbst hat keine High-Level-`inline_shapes`-API, daher
    # direkt im Roh-XML des Header-Parts nach einem eingebetteten Bild
    # gesucht (<pic:pic>-Element, von add_picture erzeugt).
    header_part = document.sections[0].header.part
    assert header_part.element.xml.count("<pic:pic") == 1


def test_export_service_embeds_signature_at_document_end(
    db_session: Session, tmp_path: Path
) -> None:
    draft = _seed_draft(db_session)
    signature_path = _write_tiny_png(tmp_path / "signature.png")
    profile = FirmProfile(
        firm_name="Kanzlei Mustermann Rechtsanwälte",
        signature_path=signature_path,
        signatory_name="Rechtsanwältin Anna Muster",
    )

    buffer = DraftDocxExportService().export_draft(draft, draft.matter, profile)

    document = DocxDocument(BytesIO(buffer.getvalue()))
    assert len(document.inline_shapes) == 1
    body_text = [p.text for p in document.paragraphs]
    assert body_text[-1] == "Rechtsanwältin Anna Muster"
    # Unterschrift steht NUR am Ende, nicht wiederholt (kein Word-Footer).
    footer_text = "".join(p.text for p in document.sections[0].footer.paragraphs)
    assert footer_text == ""


def test_export_service_skips_missing_image_files_gracefully(db_session: Session) -> None:
    """Ein in der DB referenzierter, aber von der Platte verschwundener
    Bildpfad darf den Export nicht zum Absturz bringen."""
    draft = _seed_draft(db_session)
    profile = FirmProfile(
        firm_name="Kanzlei Mustermann Rechtsanwälte",
        logo_path="/nicht/vorhanden/logo.png",
        signature_path="/nicht/vorhanden/signatur.png",
        signatory_name="Rechtsanwältin Anna Muster",
    )

    buffer = DraftDocxExportService().export_draft(draft, draft.matter, profile)

    document = DocxDocument(BytesIO(buffer.getvalue()))
    assert len(document.inline_shapes) == 0
    assert document.sections[0].header.part.element.xml.count("<pic:pic") == 0
    assert "Rechtsanwältin Anna Muster" in [p.text for p in document.paragraphs]


@pytest.fixture()
def client(db_session: Session) -> Iterator[TestClient]:
    def _override_get_db() -> Iterator[Session]:
        yield db_session

    app.dependency_overrides[get_db] = _override_get_db
    try:
        test_client = TestClient(app)
        login_as_admin(db_session, test_client)
        yield test_client
    finally:
        app.dependency_overrides.clear()


def test_export_route_returns_docx(client: TestClient, db_session: Session) -> None:
    draft = _seed_draft(db_session)

    response = client.get(f"/dashboard/drafts/{draft.id}/export.docx")

    assert response.status_code == 200
    assert response.headers["content-type"] == DOCX_MEDIA_TYPE
    assert "attachment" in response.headers["content-disposition"]
    document = DocxDocument(BytesIO(response.content))
    assert any("Erster Absatz." in p.text for p in document.paragraphs)


def test_export_route_includes_saved_firm_profile_letterhead(
    client: TestClient, db_session: Session
) -> None:
    draft = _seed_draft(db_session)
    db_session.add(FirmProfile(firm_name="Kanzlei Mustermann Rechtsanwälte"))
    db_session.commit()

    response = client.get(f"/dashboard/drafts/{draft.id}/export.docx")

    document = DocxDocument(BytesIO(response.content))
    header_text = "\n".join(p.text for p in document.sections[0].header.paragraphs)
    assert "Kanzlei Mustermann Rechtsanwälte" in header_text


def test_export_route_logs_audit_event(client: TestClient, db_session: Session) -> None:
    draft = _seed_draft(db_session)

    client.get(f"/dashboard/drafts/{draft.id}/export.docx")

    events = (
        db_session.query(AuditEvent)
        .filter_by(entity_id=draft.id, event_type="draft_exported_docx")
        .all()
    )
    assert len(events) == 1
